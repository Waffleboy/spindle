"""Word (.docx / .doc) document ingesters.

Uses python-docx for modern .docx files. For legacy .doc (OLE2 binary
format), attempts python-docx first, then falls back to extracting
readable text segments from the raw binary.
"""

import re
from pathlib import Path

import docx

from backend.config import get_settings
from backend.ingestion.common import (
    BaseIngester,
    IngestedDocument,
    register_ingester,
)


class WordIngester(BaseIngester):
    """Extracts text from Word (.docx) files."""

    def ingest(self, file_path: Path, storage_path: str) -> IngestedDocument:
        """Ingest a Word document.

        Extracts paragraph text; page_count is estimated from text length.

        Args:
            file_path: Path to the .docx file.
            storage_path: Relative storage path for reference.

        Returns:
            IngestedDocument with extracted text and empty pages list.
        """
        document = docx.Document(str(file_path))

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        chars_per_page = get_settings().word_chars_per_page
        page_count = max(1, len(text) // chars_per_page) if text else 0

        return IngestedDocument(
            original_filename=file_path.name,
            storage_path=storage_path,
            file_type="docx",
            pages=[],
            text=text,
            page_count=page_count,
        )


register_ingester("docx", WordIngester)


def _extract_text_from_binary(data: bytes) -> str:
    """Best-effort text extraction from a legacy .doc binary.

    Scans for runs of printable ASCII/Latin-1 characters (length >= 4)
    and joins them. This is intentionally simple -- it won't capture
    every character perfectly, but reliably extracts the human-readable
    body text from most OLE2 .doc files.
    """
    # Match runs of printable characters (space through tilde, common
    # punctuation, accented Latin-1 range 0xC0-0xFF).
    pattern = re.compile(rb"[\x20-\x7E\xC0-\xFF]{4,}")
    segments = [match.group().decode("latin-1") for match in pattern.finditer(data)]
    return "\n".join(segments)


class LegacyDocIngester(BaseIngester):
    """Extracts text from legacy Word (.doc) files.

    Strategy:
    1. Try opening with python-docx (works if the .doc is actually a
       renamed .docx or a transitional-format file).
    2. Fall back to extracting readable text segments from the raw binary.
    """

    def ingest(self, file_path: Path, storage_path: str) -> IngestedDocument:
        """Ingest a legacy .doc file.

        Args:
            file_path: Path to the .doc file.
            storage_path: Relative storage path for reference.

        Returns:
            IngestedDocument with extracted text and empty pages list.
        """
        # Attempt 1: try python-docx (handles renamed .docx files)
        try:
            document = docx.Document(str(file_path))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
        except Exception:
            # Attempt 2: raw binary extraction
            raw = file_path.read_bytes()
            text = _extract_text_from_binary(raw)

        if not text.strip():
            raise ValueError(
                f"Could not extract text from '{file_path.name}'. "
                "Consider converting the .doc file to .docx format."
            )

        chars_per_page = get_settings().word_chars_per_page
        page_count = max(1, len(text) // chars_per_page) if text else 0

        return IngestedDocument(
            original_filename=file_path.name,
            storage_path=storage_path,
            file_type="doc",
            pages=[],
            text=text,
            page_count=page_count,
        )


register_ingester("doc", LegacyDocIngester)

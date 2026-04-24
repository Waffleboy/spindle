"""Tests for the document ingestion layer."""

import tempfile
from pathlib import Path

import fitz  # PyMuPDF
import pytest
import xlrd
from docx import Document as DocxDocument
from openpyxl import Workbook
from PIL import Image

from backend.ingestion.common import BaseIngester, IngestedDocument, get_ingester
from backend.ingestion.csv_ingester import CsvIngester
from backend.ingestion.excel_ingester import ExcelIngester, XlsIngester
from backend.ingestion.pdf_ingester import PdfIngester
from backend.ingestion.word_ingester import LegacyDocIngester, WordIngester


# ---------------------------------------------------------------------------
# Helpers to create minimal test files
# ---------------------------------------------------------------------------

def _create_test_pdf(path: Path, text: str = "Hello World", num_pages: int = 2) -> Path:
    """Create a minimal PDF with the given text on each page."""
    doc = fitz.open()
    for i in range(num_pages):
        page = doc.new_page(width=612, height=792)
        page.insert_text((72, 72), f"{text} - Page {i + 1}")
    doc.save(str(path))
    doc.close()
    return path


def _create_test_docx(path: Path, paragraphs: list[str] | None = None) -> Path:
    """Create a minimal .docx with the given paragraphs."""
    if paragraphs is None:
        paragraphs = ["First paragraph.", "Second paragraph.", "Third paragraph."]
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path


def _create_test_xlsx(
    path: Path,
    sheets: dict[str, list[list]] | None = None,
) -> Path:
    """Create a minimal .xlsx with the given sheets and rows."""
    if sheets is None:
        sheets = {
            "Products": [["Name", "Price"], ["Widget", 9.99], ["Gadget", 19.99]],
            "Summary": [["Total", 29.98]],
        }
    wb = Workbook()
    # Remove default sheet
    wb.remove(wb.active)
    for sheet_name, rows in sheets.items():
        ws = wb.create_sheet(title=sheet_name)
        for row in rows:
            ws.append(row)
    wb.save(str(path))
    wb.close()
    return path


def _create_test_csv(path: Path, content: str | None = None) -> Path:
    """Create a minimal .csv file."""
    if content is None:
        content = "Name,Price\nWidget,9.99\nGadget,19.99\n"
    path.write_text(content, encoding="utf-8")
    return path


def _create_test_xls(
    path: Path,
    sheets: dict[str, list[list]] | None = None,
) -> Path:
    """Create a minimal .xls file using xlwt."""
    import xlwt

    if sheets is None:
        sheets = {
            "Products": [["Name", "Price"], ["Widget", 9.99], ["Gadget", 19.99]],
            "Summary": [["Total", 29.98]],
        }
    wb = xlwt.Workbook()
    for sheet_name, rows in sheets.items():
        ws = wb.add_sheet(sheet_name)
        for row_idx, row in enumerate(rows):
            for col_idx, value in enumerate(row):
                ws.write(row_idx, col_idx, value)
    wb.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Test: get_ingester registry
# ---------------------------------------------------------------------------

class TestGetIngester:
    def test_pdf_returns_pdf_ingester(self):
        ingester = get_ingester("pdf")
        assert isinstance(ingester, PdfIngester)
        assert isinstance(ingester, BaseIngester)

    def test_docx_returns_word_ingester(self):
        ingester = get_ingester("docx")
        assert isinstance(ingester, WordIngester)
        assert isinstance(ingester, BaseIngester)

    def test_xlsx_returns_excel_ingester(self):
        ingester = get_ingester("xlsx")
        assert isinstance(ingester, ExcelIngester)
        assert isinstance(ingester, BaseIngester)

    def test_csv_returns_csv_ingester(self):
        ingester = get_ingester("csv")
        assert isinstance(ingester, CsvIngester)
        assert isinstance(ingester, BaseIngester)

    def test_xls_returns_xls_ingester(self):
        ingester = get_ingester("xls")
        assert isinstance(ingester, XlsIngester)
        assert isinstance(ingester, BaseIngester)

    def test_doc_returns_legacy_doc_ingester(self):
        ingester = get_ingester("doc")
        assert isinstance(ingester, LegacyDocIngester)
        assert isinstance(ingester, BaseIngester)

    def test_unsupported_type_raises(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            get_ingester("pptx")


# ---------------------------------------------------------------------------
# Test: PDF ingestion
# ---------------------------------------------------------------------------

class TestPdfIngester:
    def test_basic_ingestion(self, tmp_path: Path):
        pdf_path = _create_test_pdf(tmp_path / "test.pdf", num_pages=3)
        ingester = PdfIngester()
        result = ingester.ingest(pdf_path, "data/originals/test.pdf")

        assert isinstance(result, IngestedDocument)
        assert result.file_type == "pdf"
        assert result.page_count == 3
        assert len(result.pages) == 3
        assert result.original_filename == "test.pdf"
        assert result.storage_path == "data/originals/test.pdf"

    def test_pages_are_pil_images(self, tmp_path: Path):
        pdf_path = _create_test_pdf(tmp_path / "test.pdf", num_pages=1)
        ingester = PdfIngester()
        result = ingester.ingest(pdf_path, "store/test.pdf")

        assert len(result.pages) == 1
        img = result.pages[0]
        assert isinstance(img, Image.Image)
        assert img.mode == "RGB"
        assert img.width > 0 and img.height > 0

    def test_text_extraction(self, tmp_path: Path):
        pdf_path = _create_test_pdf(tmp_path / "test.pdf", text="Taxonomy Test", num_pages=2)
        ingester = PdfIngester()
        result = ingester.ingest(pdf_path, "store/test.pdf")

        assert "Taxonomy Test" in result.text
        assert "Page 1" in result.text
        assert "Page 2" in result.text

    def test_metadata_contains_dpi(self, tmp_path: Path):
        pdf_path = _create_test_pdf(tmp_path / "test.pdf", num_pages=1)
        ingester = PdfIngester()
        result = ingester.ingest(pdf_path, "store/test.pdf")

        assert result.metadata.get("dpi") == 150


# ---------------------------------------------------------------------------
# Test: Word ingestion
# ---------------------------------------------------------------------------

class TestWordIngester:
    def test_basic_ingestion(self, tmp_path: Path):
        docx_path = _create_test_docx(tmp_path / "test.docx")
        ingester = WordIngester()
        result = ingester.ingest(docx_path, "data/originals/test.docx")

        assert isinstance(result, IngestedDocument)
        assert result.file_type == "docx"
        assert result.pages == []
        assert result.original_filename == "test.docx"

    def test_text_extraction(self, tmp_path: Path):
        paragraphs = ["Alpha", "Bravo", "Charlie"]
        docx_path = _create_test_docx(tmp_path / "test.docx", paragraphs=paragraphs)
        ingester = WordIngester()
        result = ingester.ingest(docx_path, "store/test.docx")

        for p in paragraphs:
            assert p in result.text

    def test_page_count_estimation(self, tmp_path: Path):
        # Short doc should be 1 page
        docx_path = _create_test_docx(tmp_path / "short.docx", paragraphs=["Short text."])
        ingester = WordIngester()
        result = ingester.ingest(docx_path, "store/short.docx")
        assert result.page_count == 1

        # Long doc should estimate more pages
        long_text = "A" * 9000  # 9000 chars -> 3 pages
        docx_path2 = _create_test_docx(tmp_path / "long.docx", paragraphs=[long_text])
        result2 = ingester.ingest(docx_path2, "store/long.docx")
        assert result2.page_count == 3

    def test_empty_paragraphs_excluded(self, tmp_path: Path):
        docx_path = _create_test_docx(
            tmp_path / "test.docx", paragraphs=["Hello", "", "  ", "World"]
        )
        ingester = WordIngester()
        result = ingester.ingest(docx_path, "store/test.docx")

        # Empty / whitespace-only paragraphs should be filtered out
        lines = result.text.split("\n")
        assert len(lines) == 2
        assert lines[0] == "Hello"
        assert lines[1] == "World"


# ---------------------------------------------------------------------------
# Test: Excel ingestion
# ---------------------------------------------------------------------------

class TestExcelIngester:
    def test_basic_ingestion(self, tmp_path: Path):
        xlsx_path = _create_test_xlsx(tmp_path / "test.xlsx")
        ingester = ExcelIngester()
        result = ingester.ingest(xlsx_path, "data/originals/test.xlsx")

        assert isinstance(result, IngestedDocument)
        assert result.file_type == "xlsx"
        assert result.pages == []
        assert result.original_filename == "test.xlsx"

    def test_sheet_count_as_page_count(self, tmp_path: Path):
        sheets = {
            "Sheet1": [["a"]],
            "Sheet2": [["b"]],
            "Sheet3": [["c"]],
        }
        xlsx_path = _create_test_xlsx(tmp_path / "test.xlsx", sheets=sheets)
        ingester = ExcelIngester()
        result = ingester.ingest(xlsx_path, "store/test.xlsx")

        assert result.page_count == 3

    def test_structured_text_format(self, tmp_path: Path):
        sheets = {
            "Data": [["Name", "Value"], ["Alpha", 1], ["Bravo", 2]],
        }
        xlsx_path = _create_test_xlsx(tmp_path / "test.xlsx", sheets=sheets)
        ingester = ExcelIngester()
        result = ingester.ingest(xlsx_path, "store/test.xlsx")

        assert "Sheet: Data" in result.text
        assert "Name\tValue" in result.text
        assert "Alpha\t1" in result.text

    def test_metadata_has_sheet_names(self, tmp_path: Path):
        sheets = {"Products": [["x"]], "Sales": [["y"]]}
        xlsx_path = _create_test_xlsx(tmp_path / "test.xlsx", sheets=sheets)
        ingester = ExcelIngester()
        result = ingester.ingest(xlsx_path, "store/test.xlsx")

        assert result.metadata["sheet_names"] == ["Products", "Sales"]


# ---------------------------------------------------------------------------
# Test: CSV ingestion
# ---------------------------------------------------------------------------

class TestCsvIngester:
    def test_basic_ingestion(self, tmp_path: Path):
        csv_path = _create_test_csv(tmp_path / "test.csv")
        ingester = CsvIngester()
        result = ingester.ingest(csv_path, "data/originals/test.csv")

        assert isinstance(result, IngestedDocument)
        assert result.file_type == "csv"
        assert result.pages == []
        assert result.page_count == 1
        assert result.original_filename == "test.csv"

    def test_text_extraction(self, tmp_path: Path):
        csv_path = _create_test_csv(tmp_path / "test.csv")
        ingester = CsvIngester()
        result = ingester.ingest(csv_path, "store/test.csv")

        assert "Name\tPrice" in result.text
        assert "Widget\t9.99" in result.text
        assert "Gadget\t19.99" in result.text

    def test_empty_rows_excluded(self, tmp_path: Path):
        content = "A,B\n,,\nX,Y\n"
        csv_path = _create_test_csv(tmp_path / "test.csv", content=content)
        ingester = CsvIngester()
        result = ingester.ingest(csv_path, "store/test.csv")

        lines = result.text.split("\n")
        assert len(lines) == 2
        assert "A\tB" in lines[0]
        assert "X\tY" in lines[1]

    def test_utf8_bom_handled(self, tmp_path: Path):
        # UTF-8 BOM prefix
        content = "﻿Name,Value\nTest,42\n"
        csv_path = tmp_path / "bom.csv"
        csv_path.write_bytes(content.encode("utf-8-sig"))
        ingester = CsvIngester()
        result = ingester.ingest(csv_path, "store/bom.csv")

        assert "Name\tValue" in result.text

    def test_latin1_fallback(self, tmp_path: Path):
        # Latin-1 encoded content with non-UTF-8 byte
        content = b"Name,Value\nCaf\xe9,1\n"
        csv_path = tmp_path / "latin.csv"
        csv_path.write_bytes(content)
        ingester = CsvIngester()
        result = ingester.ingest(csv_path, "store/latin.csv")

        assert "Caf" in result.text


class TestCsvIngesterRows:
    def test_ingest_rows_basic(self, tmp_path: Path):
        content = "Meeting,Date,Notes\nBoard Q1,2024-03-15,Revenue up\nBoard Q2,2024-06-15,Hired CEO\n"
        csv_path = _create_test_csv(tmp_path / "meetings.csv", content=content)
        ingester = CsvIngester()
        results = ingester.ingest_rows(csv_path, "store/meetings.csv")

        assert len(results) == 2
        assert results[0].original_filename == "meetings.csv [Row 1]"
        assert results[1].original_filename == "meetings.csv [Row 2]"
        assert "Meeting: Board Q1" in results[0].text
        assert "Date: 2024-03-15" in results[0].text
        assert "Notes: Revenue up" in results[0].text
        assert "Meeting: Board Q2" in results[1].text

    def test_ingest_rows_preserves_file_type(self, tmp_path: Path):
        content = "A,B\n1,2\n"
        csv_path = _create_test_csv(tmp_path / "data.csv", content=content)
        ingester = CsvIngester()
        results = ingester.ingest_rows(csv_path, "store/data.csv")

        assert len(results) == 1
        assert results[0].file_type == "csv"
        assert results[0].page_count == 1

    def test_ingest_rows_header_only_falls_back(self, tmp_path: Path):
        content = "Col1,Col2\n"
        csv_path = _create_test_csv(tmp_path / "empty.csv", content=content)
        ingester = CsvIngester()
        results = ingester.ingest_rows(csv_path, "store/empty.csv")

        assert len(results) == 1
        assert results[0].original_filename == "empty.csv"

    def test_ingest_rows_single_row(self, tmp_path: Path):
        content = "Name,Value\nAlpha,100\n"
        csv_path = _create_test_csv(tmp_path / "single.csv", content=content)
        ingester = CsvIngester()
        results = ingester.ingest_rows(csv_path, "store/single.csv")

        assert len(results) == 1
        assert results[0].original_filename == "single.csv [Row 1]"
        assert "Name: Alpha" in results[0].text
        assert "Value: 100" in results[0].text

    def test_ingest_rows_empty_cells(self, tmp_path: Path):
        content = "A,B,C\n1,,3\n"
        csv_path = _create_test_csv(tmp_path / "sparse.csv", content=content)
        ingester = CsvIngester()
        results = ingester.ingest_rows(csv_path, "store/sparse.csv")

        assert len(results) == 1
        assert "B: " in results[0].text

    def test_ingest_rows_more_cols_than_headers(self, tmp_path: Path):
        content = "A,B\n1,2,3\n"
        csv_path = _create_test_csv(tmp_path / "extra.csv", content=content)
        ingester = CsvIngester()
        results = ingester.ingest_rows(csv_path, "store/extra.csv")

        assert len(results) == 1
        assert "Column 3: 3" in results[0].text


# ---------------------------------------------------------------------------
# Test: XLS ingestion
# ---------------------------------------------------------------------------

class TestXlsIngester:
    def test_basic_ingestion(self, tmp_path: Path):
        xls_path = _create_test_xls(tmp_path / "test.xls")
        ingester = XlsIngester()
        result = ingester.ingest(xls_path, "data/originals/test.xls")

        assert isinstance(result, IngestedDocument)
        assert result.file_type == "xls"
        assert result.pages == []
        assert result.original_filename == "test.xls"

    def test_sheet_count_as_page_count(self, tmp_path: Path):
        sheets = {
            "Sheet1": [["a"]],
            "Sheet2": [["b"]],
            "Sheet3": [["c"]],
        }
        xls_path = _create_test_xls(tmp_path / "test.xls", sheets=sheets)
        ingester = XlsIngester()
        result = ingester.ingest(xls_path, "store/test.xls")

        assert result.page_count == 3

    def test_structured_text_format(self, tmp_path: Path):
        sheets = {
            "Data": [["Name", "Value"], ["Alpha", 1], ["Bravo", 2]],
        }
        xls_path = _create_test_xls(tmp_path / "test.xls", sheets=sheets)
        ingester = XlsIngester()
        result = ingester.ingest(xls_path, "store/test.xls")

        assert "Sheet: Data" in result.text
        assert "Name\tValue" in result.text
        assert "Alpha\t1" in result.text

    def test_metadata_has_sheet_names(self, tmp_path: Path):
        sheets = {"Products": [["x"]], "Sales": [["y"]]}
        xls_path = _create_test_xls(tmp_path / "test.xls", sheets=sheets)
        ingester = XlsIngester()
        result = ingester.ingest(xls_path, "store/test.xls")

        assert result.metadata["sheet_names"] == ["Products", "Sales"]

    def test_integer_rendering(self, tmp_path: Path):
        """Integers should not have trailing .0."""
        sheets = {"Data": [["Count"], [42]]}
        xls_path = _create_test_xls(tmp_path / "test.xls", sheets=sheets)
        ingester = XlsIngester()
        result = ingester.ingest(xls_path, "store/test.xls")

        assert "42" in result.text
        assert "42.0" not in result.text


# ---------------------------------------------------------------------------
# Test: Legacy .doc ingestion
# ---------------------------------------------------------------------------

class TestLegacyDocIngester:
    def test_renamed_docx_file(self, tmp_path: Path):
        """A .docx file renamed to .doc should be handled by python-docx."""
        docx_path = _create_test_docx(tmp_path / "test.docx", paragraphs=["Hello legacy"])
        # Rename to .doc
        doc_path = tmp_path / "test.doc"
        docx_path.rename(doc_path)

        ingester = LegacyDocIngester()
        result = ingester.ingest(doc_path, "store/test.doc")

        assert isinstance(result, IngestedDocument)
        assert result.file_type == "doc"
        assert "Hello legacy" in result.text
        assert result.pages == []
        assert result.page_count >= 1

    def test_binary_fallback_extracts_text(self, tmp_path: Path):
        """A binary file with readable text segments should have text extracted."""
        # Create a fake binary .doc with embedded readable text
        binary_content = (
            b"\x00\x00\x00"
            b"This is readable text from a legacy document"
            b"\x00\x01\x02\x03"
            b"Another readable segment here"
            b"\xff\xfe\x00"
        )
        doc_path = tmp_path / "legacy.doc"
        doc_path.write_bytes(binary_content)

        ingester = LegacyDocIngester()
        result = ingester.ingest(doc_path, "store/legacy.doc")

        assert "This is readable text from a legacy document" in result.text
        assert "Another readable segment here" in result.text

    def test_empty_binary_raises(self, tmp_path: Path):
        """A binary file with no extractable text should raise ValueError."""
        doc_path = tmp_path / "empty.doc"
        doc_path.write_bytes(b"\x00\x01\x02\x03\x04\x05")

        ingester = LegacyDocIngester()
        with pytest.raises(ValueError, match="Could not extract text"):
            ingester.ingest(doc_path, "store/empty.doc")


# ---------------------------------------------------------------------------
# Test: store_and_ingest service
# ---------------------------------------------------------------------------

def _patch_settings(monkeypatch, tmp_path):
    """Override get_settings to redirect originals_dir to tmp_path."""
    from backend.config import Settings

    test_settings = Settings(originals_dir=tmp_path, data_dir=tmp_path)
    _getter = lambda: test_settings  # noqa: E731
    monkeypatch.setattr("backend.config.get_settings", _getter)
    monkeypatch.setattr("backend.ingestion.service.get_settings", _getter)
    monkeypatch.setattr("backend.ingestion.pdf_ingester.get_settings", _getter)
    monkeypatch.setattr("backend.ingestion.word_ingester.get_settings", _getter)


class TestStoreAndIngest:
    def test_pdf_store_and_ingest(self, tmp_path: Path, monkeypatch):
        """Test that store_and_ingest writes to disk, creates DB record, and ingests."""
        _patch_settings(monkeypatch, tmp_path)

        # Use an in-memory SQLite DB for isolation
        from sqlalchemy import create_engine
        from sqlalchemy.orm import sessionmaker

        from backend.database import Base

        engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False})
        Base.metadata.create_all(bind=engine)
        TestSession = sessionmaker(bind=engine)
        monkeypatch.setattr("backend.ingestion.service.SessionLocal", TestSession)

        # Create a test PDF in memory
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((72, 72), "Service test content")
        pdf_bytes = pdf_doc.tobytes()
        pdf_doc.close()

        from backend.ingestion.service import store_and_ingest

        doc, ingested = store_and_ingest("report.pdf", pdf_bytes)

        # Check DB record
        assert doc.original_filename == "report.pdf"
        assert doc.file_type == "pdf"
        assert doc.page_count == 1
        assert doc.id is not None

        # Check IngestedDocument
        assert ingested.file_type == "pdf"
        assert ingested.page_count == 1
        assert "Service test content" in ingested.text
        assert len(ingested.pages) == 1

        # Check file was written to disk
        stored_file = tmp_path / Path(ingested.storage_path).name
        assert stored_file.exists()
        assert stored_file.stat().st_size > 0

    def test_docx_store_and_ingest(self, tmp_path: Path, monkeypatch):
        _patch_settings(monkeypatch, tmp_path)

        from sqlalchemy import create_engine
        from sqlalchemy.orm import sessionmaker

        from backend.database import Base

        engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False})
        Base.metadata.create_all(bind=engine)
        TestSession = sessionmaker(bind=engine)
        monkeypatch.setattr("backend.ingestion.service.SessionLocal", TestSession)

        # Create docx bytes
        import io

        docx_doc = DocxDocument()
        docx_doc.add_paragraph("Test paragraph")
        buf = io.BytesIO()
        docx_doc.save(buf)
        docx_bytes = buf.getvalue()

        from backend.ingestion.service import store_and_ingest

        doc, ingested = store_and_ingest("notes.docx", docx_bytes)

        assert doc.file_type == "docx"
        assert ingested.file_type == "docx"
        assert "Test paragraph" in ingested.text

    def test_unsupported_extension_raises(self, tmp_path: Path, monkeypatch):
        _patch_settings(monkeypatch, tmp_path)

        from backend.ingestion.service import store_and_ingest

        with pytest.raises(ValueError, match="Unsupported file extension"):
            store_and_ingest("data.pptx", b"dummy")


class TestStoreAndIngestCsvRows:
    def test_csv_rows_creates_multiple_documents(self, tmp_path: Path, monkeypatch):
        _patch_settings(monkeypatch, tmp_path)

        from sqlalchemy import create_engine
        from sqlalchemy.orm import sessionmaker

        from backend.database import Base
        from backend.models import Document

        engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False})
        Base.metadata.create_all(bind=engine)
        TestSession = sessionmaker(bind=engine)
        monkeypatch.setattr("backend.ingestion.service.SessionLocal", TestSession)

        csv_content = b"Meeting,Date\nBoard Q1,2024-03-15\nBoard Q2,2024-06-15\n"

        from backend.ingestion.service import store_and_ingest_csv_rows

        results = store_and_ingest_csv_rows("meetings.csv", csv_content)

        assert len(results) == 2
        doc1, ingested1 = results[0]
        doc2, ingested2 = results[1]

        assert doc1.original_filename == "meetings.csv [Row 1]"
        assert doc2.original_filename == "meetings.csv [Row 2]"
        assert doc1.file_type == "csv"
        assert doc1.id != doc2.id

        assert "Meeting: Board Q1" in ingested1.text
        assert "Meeting: Board Q2" in ingested2.text

        session = TestSession()
        db_docs = session.query(Document).all()
        assert len(db_docs) == 2
        session.close()
